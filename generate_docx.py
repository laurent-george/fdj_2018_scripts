"""
Un script pour convertir un csv en document .doc avec des styles
"""


import csv

from collections import OrderedDict, defaultdict

from docxtpl import DocxTemplate
from docx.document import Document


input_csv = 'test_fdj_2018.csv'
template_docx = 'programme_example_tpl.docx'


def merge_column_with_same_name(input_csv):
    """
    Fonction permetant de gérer les colonnes identique comme Choix du jardin
    :param input_csv:
    :return:
    """
    # merge column with same name
    with open(input_csv, 'r') as f:
        basic_reader = csv.reader(f, delimiter=',')
        header = next(basic_reader)
        res = []
        for entry in basic_reader:
            cur_entry = OrderedDict()
            for key, val in zip(header, entry):
                if not key in cur_entry or val != '':
                    cur_entry[key] = val
            res.append(cur_entry)
    return res



def preprocess_csv(input_csv, output_csv):
    """
    preprocess the input csv in order to get something sorted/more easy to process
    :param input_csv:
    :param output_csv:
    :return:
    """
    res = merge_column_with_same_name(input_csv)

    if len(res) == 0:
        return


    # sorting
    res = sorted(res, key=lambda x: (x['Emplacement du jardin'],
                                     x['Choix du jardin']))  # on trie par arrondissement puis par choix du jardin

    # write new csv with unique column name
    with open(output_csv, 'w') as f:
        basic_writer = csv.DictWriter(f, delimiter=',', fieldnames=res[0].keys())
        basic_writer.writeheader()
        for row in res:
            basic_writer.writerow(row)


def _create_new_entry_in_doc(row, document):

    document.add_paragraph(row['Choix du jardin'], 'nom du lieu')
    document.add_paragraph(row["Description de l'animation"], 'liste')
    # TODO: ajouter les horaires
    document.add_paragraph("{num} {street}".format(num=row['Numéro'], street=row['Nom de la voie']),
                           style='adresse')
    document.add_paragraph("{metro}".format(metro=row['Station de métro ou RER']), style='adresse')
    details = row["Détails (ex. : face au n°12, au fond de l'allée, etc.)"]
    if details:
        document.add_paragraph("{details}".format(details=details), style='adresse')


def _add_localisation_header(localisation, document: Document):
    """
    Add a seperation for localisation
    :return:
    """
    document.add_paragraph(localisation, style='Heading1')


def main():

    tmp_file = 'merge_columnout.csv'
    preprocess_csv(input_csv, tmp_file)

    tpl=DocxTemplate(template_docx)
    sd = tpl.new_subdoc()
    prev = None

    with open(tmp_file, 'r') as f:
        reader = csv.DictReader(f, delimiter=',')
        for row in reader:
            if prev is None or prev['Emplacement du jardin'] != row['Emplacement du jardin']:
                _add_localisation_header(row['Emplacement du jardin'], document=sd)
            _create_new_entry_in_doc(row, document=sd)
            prev = row

    context = {
        'mysubdoc' : sd,
    }

    tpl.render(context)
    tpl.save('output.docx')


if __name__ == "__main__":
    main()
